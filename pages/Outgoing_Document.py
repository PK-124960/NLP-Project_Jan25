import streamlit as st
import numpy as np
import json
import requests
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer
from io import BytesIO
from docx import Document
from docx.shared import Pt

# ---------- CONFIG ----------
st.set_page_config(page_title="สร้างหนังสือออก (Outgoing Letter)", layout="wide")

# ---------- CUSTOM EARTH THEME ----------
st.markdown("""
    <style>
        [data-testid="stAppViewContainer"] > .main {
            background-color: #F5F5DC;
        }
        .earth-card {
            background-color: #FAF0E6;
            padding: 20px;
            border-radius: 12px;
            border: 2px solid #D2B48C;
            margin-bottom: 20px;
        }
        h3, h2, h1 {
            color: #5B4636;
        }
        .stButton>button {
            background-color: #8FBC8F;
            color: white;
            border-radius: 10px;
            padding: 10px 20px;
            font-weight: bold;
        }
        .stButton>button:hover {
            background-color: #6B8E23;
        }
    </style>
""", unsafe_allow_html=True)

# ---------- FUNCTIONS for Section 1 ----------
def create_prompt_fewshot(input_sec2, input_sec1):
    prompt = f"""
คุณคือผู้เชี่ยวชาญสรุปหนังสือราชการทางทหาร
ให้อ่านข้อมูลที่กำหนดอย่างละเอียด และสรุปเนื้อหาอย่างกระชับ ชัดเจน และเป็นภาษาทางการ
ห้ามแต่งข้อมูลเพิ่มนอกเหนือจากที่กำหนด
ให้นำเสนอเป็นข้อความย่อความเท่านั้น ไม่ต้องจัดหมวดหมู่ ไม่ต้องตอบเป็น JSON

เนื้อหาสรุปหนังสือราชการ ด้วยบริบท ใคร ทำอะไร ที่ไหน เมื่อไร วิเคราะห์จากข้อความนี้:
\"\"\"{input_sec2}\"\"\"

เหตุผลเบื้องหลัง:
\"\"\"{input_sec1}\"\"\"

โปรดตอบเฉพาะสรุปข้อความย่อเป็นภาษาไทย จำนวน 1 ย่อหน้าเท่านั้น 
ห้ามใส่ข้อความลักษณะ ผู้ช่วยสรุป : หรือ เหตุผลเบื้องหลัง :
ห้ามใส่หมายเลขลำดับข้อ
"""
    return prompt

def abstractive_summarize(input_sec2, input_sec1, temperature=0.2):
    url = "http://localhost:11434/api/generate"
    prompt = create_prompt_fewshot(input_sec2, input_sec1)

    payload = {
        "model": "wangchanglm",
        "prompt": prompt,
        "temperature": temperature,
        "stream": True
    }
    response = requests.post(url, json=payload, stream=True)

    collected_response = ""
    for line in response.iter_lines():
        if line:
            decoded_line = line.decode('utf-8')
            data = json.loads(decoded_line)
            if 'response' in data:
                collected_response += data['response']

    return collected_response.strip()

def calculate_cosine_similarity(text1, text2):
    emb1 = similarity_model.encode(text1)
    emb2 = similarity_model.encode(text2)
    score = cosine_similarity([emb1], [emb2])[0][0]
    return score

# ---------- FUNCTIONS for Section 2 ----------

def create_prompt_for_fact(input_sec1):
    prompt = f"""
คุณคือผู้ช่วยร่างหนังสือราชการทางทหาร
ให้อ่านข้อมูล "ข้อ ๑" ด้านล่างอย่างละเอียด
และเขียน "ข้อ ๒" ซึ่งเป็นการบรรยายข้อเท็จจริงเพิ่มเติม
โดยมีแนวทางดังนี้:

- สรุปบริบททางทหาร เช่น กำลังพล หน่วยงานที่เกี่ยวข้อง สถานที่และเวลา
- ระบุบทบาทหน้าที่ของหน่วยที่กล่าวถึง ว่ามีความเกี่ยวข้องอย่างไร
- เนื้อหาต้องเป็นภาษาราชการ กะทัดรัด ชัดเจน ไม่เยิ่นเย้อ
- ตอบเป็นข้อความย่อหน้าเดียว (ไม่แบ่งข้อย่อย)
- สามารถแต่งเติมข้อมูลที่เหมาะสมเพื่อความสมจริงได้ เช่น จำนวนผู้เข้าร่วม ภารกิจหลัก เพื่อการสาธิต
- ห้ามแต่งเรื่องนอกเหนือบริบทของข้อ ๑ อย่างสิ้นเชิง
- เป็นภาษาราชการ

เนื้อหาข้อ ๑:
\"\"\"{input_sec1}\"\"\"

โปรดร่างข้อความข้อ ๒ ตามหลักเกณฑ์ข้างต้น
"""
    return prompt

def generate_fact_from_sec1(input_sec1):
    url = "http://localhost:11434/api/generate"
    prompt = create_prompt_for_fact(input_sec1)

    payload = {
        "model": "wangchanglm",
        "prompt": prompt,
        "temperature": 0.2,
        "stream": True
    }
    response = requests.post(url, json=payload, stream=True)

    collected_response = ""
    for line in response.iter_lines():
        if line:
            decoded_line = line.decode('utf-8')
            data = json.loads(decoded_line)
            if 'response' in data:
                collected_response += data['response']

    return collected_response.strip()

# ---------- FUNCTIONS for Section 3 ----------

def create_prompt_for_section3(input_sec1, input_sec2, with_feedback=False):
    base_prompt = f"""
<fact>
ข้อ ๑: {input_sec1}

ข้อ ๒: {input_sec2}
</fact>

<instruction>
โปรดเขียนข้อเสนอหรือสั่งการ (ข้อ ๓) โดยปฏิบัติตามแนวทางต่อไปนี้:
- ลำดับเป็นข้อ ๆ: 3.1, 3.2, 3.3
- ห้ามเกิน 3 ข้อ
- แต่ละข้อเป็นประโยคคำสั่งสั้น ๆ (1–2 บรรทัด)
- ใช้ภาษาราชการที่สุภาพ และสอดคล้องกับข้อ ๑ และข้อ ๒
- ห้ามมีข้อความสรุปท้าย
</instruction>
"""

    if with_feedback:
        feedback_example = """
<example>
2.1 เรียนเชิญ ผบ.รร.ชท. หรือผู้แทนเข้าร่วมพิธีเปิดฯ ตามข้อ 1
2.2 กสน.ฯ จัดรถรับ-ส่ง เข้าร่วมพิธีเปิดฯ
2.3 ผธก.ฯ บันทึกลงระบบสารบรรณอิเล็กทรอนิกส์ (ECM) ให้หน่วยที่เกี่ยวข้องดำเนินการต่อไป
</example>
"""
        base_prompt = feedback_example.strip() + "\n" + base_prompt.strip()

    return base_prompt.strip()

def generate_section3(input_sec1, input_sec2, with_feedback=False):
    url = "http://localhost:11434/api/generate"

    prompt = create_prompt_for_section3(input_sec1, input_sec2, with_feedback=with_feedback)

    payload = {
        "model": "wangchanglm",
        "prompt": prompt,
        "temperature": 0.2,
        "stream": True
    }
    response = requests.post(url, json=payload, stream=True)

    collected_response = ""
    for line in response.iter_lines():
        if line:
            decoded_line = line.decode('utf-8')
            data = json.loads(decoded_line)
            if 'response' in data:
                collected_response += data['response']

    return collected_response.strip()

def regenerate_section3_with_feedback():
    # ดึงข้อมูลจาก session state
    sec1 = st.session_state.get("outgoing_section1", "")
    sec2 = st.session_state.get("outgoing_section2", "")
    
    if not sec1 or not sec2:
        st.error("❌ ไม่พบข้อมูลข้อ ๑ หรือข้อมูลประกอบ กรุณาสร้างข้อ ๑ ก่อน")
        return ""

    # เรียก generate ใหม่
    new_section3 = generate_section3(sec1, sec2, with_feedback=True)
    return new_section3


# ---------- FUNCTIONS ----------
def generate_docx(section1, section2, section3):
    doc = Document()

    # 📜 ตั้งค่าฟอนต์เริ่มต้น
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH SarabunPSK'
    font.size = Pt(16)

    doc.add_heading('บันทึกข้อความ', 0)

    # 📜 ข้อ ๑
    doc.add_paragraph(f'๑. {section1}', style='Normal')
    doc.add_paragraph()

    # 📜 ข้อ ๒
    doc.add_paragraph(f'๒. {section2}', style='Normal')
    doc.add_paragraph()

    # 📜 ข้อ ๓
    doc.add_paragraph(f'๓. {section3}', style='Normal')
    doc.add_paragraph()

    # 📜 ลายเซ็น (พื้นที่เผื่อไว้)
    doc.add_paragraph("\n(ลงชื่อ) ................................................")
    doc.add_paragraph("        (.............................................)  ")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------- LOAD EMBEDDER ----------
similarity_model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')

# ---------- UI START ----------
st.title("📄 สร้างหนังสือส่ง (Outgoing Letter Creation)")

# 🔹 ตัวอย่างอ้างอิง
reference_output = "ยบ.ทหาร (สนพ.ยบ.ทหาร) กําหนดพิธีเปิดการฝึกอบรมหลักสูตรนายทหารประทวน สายวิทยาการแพทย์ รุ่นที่ 3 ประจําปีงบประมาณ พ.ศ. 2568 ในวันพุธที่ 5 พ.ย. 67 เวลา 1300 ณ ห้องประชุม หทัยนเรศ ชั้น 2 สนพ.ยบ.ทหาร (บางซ่อน) โดยมี จก.ยบ.ทหาร เป็นประธานฯ การแต่งกาย เครื่องแบบปกติ คอพับแขนยาว (ทอ. อินทรธนูแข็ง) งดหมวก รายละเอียดตามสิ่งที่ส่งมาด้วย "

if "corrected_sections" not in st.session_state:
    st.error("❌ ไม่พบข้อมูลจากหนังสือรับ")
else:
    corrected_data = st.session_state["corrected_sections"]
    input_section1 = corrected_data.get("ข้อ 1", "").strip()
    input_section2 = corrected_data.get("ข้อ 2", "").strip()

    if not input_section1:
        st.error("❌ ไม่มีข้อมูลข้อ ๑")
    else:
        st.markdown("<div class='earth-card'>", unsafe_allow_html=True)
        st.subheader("🔎 ขั้นตอนสรุปเนื้อหาข้อ ๑ (Sec.1 creation)")

        temperature_values = [0.1, 0.2, 0.3, 0.4, 0.5]
        summaries = []
        cosine_scores = []

        with st.spinner('🧠 กำลังประมวลผล LLM หลายเวอร์ชัน...'):
            for temp in temperature_values:
                summary = abstractive_summarize(input_section2, input_section1, temperature=temp)
                score = calculate_cosine_similarity(reference_output, summary)
                summaries.append(summary)
                cosine_scores.append(score)

        best_idx = np.argmax(cosine_scores)

        for idx, (summary, score) in enumerate(zip(summaries, cosine_scores)):
            color = "#e0ffe0" if idx == best_idx else "#f0f0f0"
            with st.expander(f"📑 เวอร์ชัน {idx+1} (Temp {temperature_values[idx]} / Cosine {score:.4f})", expanded=(idx==best_idx)):
                st.markdown(f"<div style='background-color:{color}; padding:15px; border-radius:10px;'>{summary}</div>", unsafe_allow_html=True)

        options = [f"เวอร์ชัน {i+1} (temp={temperature_values[i]}, similarity={cosine_scores[i]:.4f})" for i in range(len(summaries))]
        default_option = options[best_idx]  # ต้องเป็น String เท่านั้น
        selected = st.radio("เลือกเวอร์ชัน (Select best version)", options, index=options.index(default_option))

        selected_summary = summaries[options.index(selected)]

        edited_section1 = st.text_area("📋 แก้ไขข้อความข้อ ๑ (customize for sec.1)", value=selected_summary, height=300)
        st.markdown("</div>", unsafe_allow_html=True)

    # ------------------ 1. บันทึกข้อ 1 และสร้างข้อ 2 ทันที ------------------
    if st.button("✅ ยืนยันข้อ ๑ (Confirm Sec.1)"):
        if edited_section1.strip():
            # Save ข้อ 1
            st.session_state["outgoing_section1"] = edited_section1

            # Generate ข้อ 2 จาก ข้อ 1 ทันที
            with st.spinner('✏️ กำลังสร้างข้อเท็จจริง ข้อ ๒...'):
                st.session_state["generated_section2"] = generate_fact_from_sec1(edited_section1)

            st.session_state["confirmed_sec1"] = True
            st.success("✅ ยืนยันข้อ ๑ และสร้างข้อ ๒ เรียบร้อย กรุณาตรวจสอบข้อ ๒")

    # ------------------ 2. แสดงข้อ 2 ให้แก้ไขและยืนยัน ------------------
    if st.session_state.get("confirmed_sec1") and "generated_section2" in st.session_state:
        st.subheader("📋 ข้อ ๒: ตรวจสอบและแก้ไขข้อความ")

        edited_section2 = st.text_area(
            "📋 แก้ไขข้อความข้อ ๒ (Customize Sec.2)", 
            value=st.session_state["generated_section2"],
            height=300,
            key="section2_editor"
        )

        if st.button("✅ ยืนยันข้อ ๒ (Confirm Sec.2)"):
            st.session_state["outgoing_section2"] = edited_section2
            st.session_state["confirmed_sec2"] = True
            st.success("✅ ยืนยันข้อ ๒ เรียบร้อย กรุณาสร้างข้อ ๓ ต่อไป")

    # ------------------ 3. สร้างข้อ 3 เมื่อยืนยันข้อ 2 ------------------
    if st.session_state.get("confirmed_sec2"):
        st.subheader("✍️ ขั้นตอนสร้างข้อ ๓ (Sec.3 Creation)")

        with st.spinner('🧠 กำลังสร้างข้อ ๓...'):
            sec1 = st.session_state["outgoing_section1"]
            sec2 = st.session_state["outgoing_section2"]
            sec3_with_feedback = generate_section3(sec1, sec2, with_feedback=True)
            sec3_without_feedback = generate_section3(sec1, sec2, with_feedback=False)

        choice_sec3 = st.radio("เลือกรูปแบบการสร้างข้อ ๓", ["มี Feedback", "ไม่มี Feedback"], index=0)

        if choice_sec3 == "มี Feedback":
            if "edited_sec3_feedback" not in st.session_state:
                st.session_state["edited_sec3_feedback"] = sec3_with_feedback
            selected_sec3 = st.session_state["edited_sec3_feedback"]
        else:
            selected_sec3 = sec3_without_feedback

        edited_section3 = st.text_area(
            "📋 แก้ไขข้อความข้อ ๓ (Customize Sec.3)", 
            value=selected_sec3,
            height=300,
            key="section3_editor"
        )

        # ปุ่ม Regenerate สำหรับ Feedback เท่านั้น
        if choice_sec3 == "มี Feedback":
            if st.button("🔄 Regenerate ข้อ ๓ (มี Feedback)"):
                new_sec3 = regenerate_section3_with_feedback()
                st.session_state["edited_sec3_feedback"] = new_sec3
                st.rerun()

        if st.button("✅ ยืนยันข้อ ๓ (Confirm Sec.3)", key="confirm_sec3_button"):
            st.session_state["outgoing_section3"] = edited_section3
            st.session_state["confirmed_sec3"] = True
            st.success("✅ ยืนยันข้อ ๓ เรียบร้อย กรุณา Preview เอกสาร")

    # ------------------ 4. Preview และดาวน์โหลด เมื่อยืนยันครบทั้ง 3 ข้อ ------------------
    if all(st.session_state.get(k) for k in ["confirmed_sec1", "confirmed_sec2", "confirmed_sec3"]):
        st.markdown("<div class='earth-card'>", unsafe_allow_html=True)
        st.subheader("📋 Preview หนังสือส่ง")

        kru_logo = "kru_logo.png"
        st.image(kru_logo, width=100)

        st.markdown("<h3 style='text-align: center;'>บันทึกข้อความ</h3>", unsafe_allow_html=True)

        corrected_data = st.session_state.get("corrected_sections", {})
        st.markdown(f"""
        **ส่วนราชการ**: {corrected_data.get('ส่วนราชการ', 'ไม่มีข้อมูล')}  
        **ที่**: {corrected_data.get('ที่', 'ไม่มีข้อมูล')}  
        **วันที่**: {corrected_data.get('วันที่', 'ไม่มีข้อมูล')}  
        **เรื่อง**: {corrected_data.get('เรื่อง', 'ไม่มีข้อมูล')}  
        **เรียน**: {corrected_data.get('เรียน', 'ไม่มีข้อมูล')}
        """, unsafe_allow_html=True)

        if corrected_data.get("อ้างถึง"):
            st.markdown(f"**อ้างถึง**: {corrected_data.get('อ้างถึง')}")
        if corrected_data.get("สิ่งที่ส่งมาด้วย"):
            st.markdown(f"**สิ่งที่ส่งมาด้วย**: {corrected_data.get('สิ่งที่ส่งมาด้วย')}")

        st.markdown("<br>", unsafe_allow_html=True)

        section_contents = [
            ("๑", f"ตามที่ {st.session_state['outgoing_section1']}"),
            ("๒", st.session_state["outgoing_section2"]),
            ("๓", f"เพื่อให้การดำเนินการเป็นไปด้วยความเรียบร้อย {st.session_state['outgoing_section3']}")
        ]

        for thai_number, display_content in section_contents:
            st.markdown(f"""
            <div style='text-indent: 2em; font-size: 1.1rem; line-height: 1.8;'>
                <strong>{thai_number}.</strong> {display_content}
            </div>
            <br>
            """, unsafe_allow_html=True)

        st.markdown("""
        <p style='text-align: right;'>
            (ลงชื่อ) ................................................<br>
            (.............................................)
        </p>
        """, unsafe_allow_html=True)

        st.success("✅ เอกสารแสดงผลเรียบร้อยแล้ว!")

        st.markdown("</div>", unsafe_allow_html=True)

        docx_file = generate_docx(
            st.session_state["outgoing_section1"],
            st.session_state["outgoing_section2"],
            st.session_state["outgoing_section3"]
        )

        st.download_button(
            "📥 ดาวน์โหลด DOCX",
            data=docx_file,
            file_name="Outgoing_Letter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

